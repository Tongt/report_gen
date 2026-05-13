from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from pypdf import PdfReader

from config import KB_DIR, SKILLS_DIR, SUPPORTED_EXTENSIONS, TEMPLATES_DIR


@dataclass
class LoadedDocument:
    file_name: str
    source_path: str
    text: str


KB_FACT_EXTENSIONS = {".pdf", ".docx", ".md"}
EXCLUDED_DIR_NAMES = {"skills", "templates", "prompts", "chroma_db", "__pycache__"}
EXCLUDED_FILE_NAMES = {"README.md", "requirements.txt", ".env", "app.py"}


def save_uploaded_files(uploaded_files: Iterable, target_dir: Path) -> list[Path]:
    target_dir.mkdir(parents=True, exist_ok=True)
    saved_paths: list[Path] = []
    for upload in uploaded_files:
        out_path = target_dir / upload.name
        if out_path.exists():
            continue
        out_path.write_bytes(upload.getbuffer())
        saved_paths.append(out_path)
    return saved_paths


def read_file_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return ""
    if suffix == ".txt" or suffix == ".md":
        return file_path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        reader = PdfReader(str(file_path))
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        return "\n".join(pages)
    if suffix == ".docx":
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    return ""


def load_documents(file_paths: list[Path]) -> list[LoadedDocument]:
    docs: list[LoadedDocument] = []
    for path in file_paths:
        text = read_file_text(path).strip()
        if not text:
            continue
        docs.append(
            LoadedDocument(
                file_name=path.name,
                source_path=str(path),
                text=text,
            )
        )
    return docs


def _is_allowed_knowledge_file(path: Path, base_dir: Path) -> bool:
    if not path.is_file():
        return False
    if path.suffix.lower() not in KB_FACT_EXTENSIONS:
        return False
    if path.name.endswith(".skill.md"):
        return False
    if path.name in EXCLUDED_FILE_NAMES:
        return False
    try:
        relative_parts = set(path.relative_to(base_dir).parts)
    except Exception:
        relative_parts = set(path.parts)
    if relative_parts & EXCLUDED_DIR_NAMES:
        return False
    return True


def list_knowledge_base_files(base_dir: Path = KB_DIR) -> list[Path]:
    base_dir.mkdir(parents=True, exist_ok=True)
    files = [
        path
        for path in base_dir.rglob("*")
        if _is_allowed_knowledge_file(path, base_dir=base_dir)
    ]
    return sorted(files, key=lambda p: str(p).lower())


def load_knowledge_base(base_dir: Path = KB_DIR) -> list[LoadedDocument]:
    return load_documents(list_knowledge_base_files(base_dir=base_dir))


def load_skills(skills_dir: Path = SKILLS_DIR) -> dict[str, str]:
    skills_dir.mkdir(parents=True, exist_ok=True)
    output: dict[str, str] = {}
    for path in sorted(skills_dir.glob("*.skill.md"), key=lambda p: p.name.lower()):
        text = path.read_text(encoding="utf-8", errors="ignore").strip()
        if text:
            output[path.name] = text
    return output


def load_templates(templates_dir: Path = TEMPLATES_DIR) -> dict[str, str]:
    templates_dir.mkdir(parents=True, exist_ok=True)
    output: dict[str, str] = {}
    for path in sorted(templates_dir.glob("*.md"), key=lambda p: p.name.lower()):
        if path.name.endswith(".skill.md"):
            continue
        text = path.read_text(encoding="utf-8", errors="ignore").strip()
        if text:
            output[path.name] = text
    return output
